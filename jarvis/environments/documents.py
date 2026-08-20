"""Document environment — read what is inside files, not just their names.

PDF, Word, Excel, CSV and plain text. This is the difference between "I found
budget.xlsx" and "budget.xlsx says you spent 40,000 on equipment".

Design decisions worth recording:

* Read-only. Writing documents is a separate, riskier capability: a malformed
  write silently corrupts something the user cannot easily reconstruct, and the
  verification for "did I write this correctly" is much harder than for "did I
  read this correctly".

* Text is extracted with a hard character budget. A 400-page PDF will not fit
  in a model's context, and truncating at the end is better than failing - but
  the truncation is always reported, never silent, because an answer drawn from
  the first 3% of a document while implying it read all of it is a lie.

* Excel is read with values, not formulas (`data_only=True`). A user asking
  what a spreadsheet says means the numbers, not "=SUM(B2:B40)". Files never
  opened in Excel have no cached values, and that is reported rather than
  returning a sheet of None.

* Encrypted PDFs are detected and reported. pypdf raises a generic error
  otherwise, which surfaces as an unhelpful failure.
"""

from __future__ import annotations

import csv
import time
from pathlib import Path
from typing import Any

from jarvis.core.contracts import ActionResult, VerificationResult
from jarvis.policy import guardrails

MAX_CHARS = 24_000
MAX_ROWS = 400
# Anything that is genuinely plain text. Kept explicit rather than "try to
# decode and see": a wrong guess on a binary file returns a screen of
# replacement characters, which is worse than refusing.
TEXT_SUFFIXES = {
    # documents and notes
    ".txt", ".md", ".rst", ".tex", ".bib", ".log",
    # data and config
    ".json", ".xml", ".yml", ".yaml", ".toml", ".ini", ".cfg", ".conf", ".env",
    # code
    ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".css", ".scss",
    ".c", ".h", ".cpp", ".hpp", ".cs", ".java", ".go", ".rs", ".rb", ".php",
    ".sql", ".r", ".m", ".swift", ".kt", ".lua", ".pl",
    # shell
    ".sh", ".bash", ".ps1", ".bat", ".cmd",
}


def _expand(raw: str) -> Path:
    import os
    text = str(raw).strip().strip('"').strip("'")
    text = os.path.expandvars(text)
    p = Path(text).expanduser()
    if not p.is_absolute():
        low = text.replace("\\", "/").lower()
        for name in ("desktop", "documents", "downloads", "pictures"):
            if low.startswith(name + "/") or low == name:
                rest = text.replace("\\", "/")[len(name):].lstrip("/")
                base = Path.home() / name.capitalize()
                return base / rest if rest else base
        p = Path.cwd() / p
    return p


def _truncate(text: str) -> tuple[str, bool]:
    if len(text) <= MAX_CHARS:
        return text, False
    return text[:MAX_CHARS], True


class DocumentEnvironment:
    """Read the contents of documents."""

    id = "documents"

    def state(self) -> dict[str, Any]:
        return {"available": True, "formats": sorted(self.formats())}

    @staticmethod
    def formats() -> set[str]:
        """Which extensions can actually be read on this machine right now."""
        supported = {".csv", ".tsv"} | TEXT_SUFFIXES
        for suffix, module in ((".pdf", "pypdf"), (".docx", "docx"),
                               (".xlsx", "openpyxl")):
            try:
                __import__(module)
                supported.add(suffix)
            except ImportError:
                pass
        if ".xlsx" in supported:
            supported.add(".xlsm")
        return supported

    def capabilities(self) -> list[str]:
        return ["read_document", "document_info"]

    def constraints(self) -> list[str]:
        return [
            "Read-only: documents are never written or modified here.",
            f"Text is capped at {MAX_CHARS} characters; truncation is always "
            f"reported rather than hidden.",
            "Spreadsheets are read as values, not formulas; a file never "
            "opened in Excel may have no cached values to read.",
            "Scanned PDFs contain images, not text - OCR is not implemented, "
            "and such a file is reported as having no extractable text.",
            "Encrypted PDFs cannot be opened without the password.",
        ]

    def act(self, ability_id: str, args: dict[str, Any]) -> ActionResult:
        handlers = {"read_document": self._read, "document_info": self._info}
        handler = handlers.get(ability_id)
        if handler is None:
            return ActionResult(ok=False, error="unregistered",
                                summary=f"unknown ability '{ability_id}'")
        raw = args.get("path") or args.get("file") or ""
        if not str(raw).strip():
            return ActionResult(ok=False, error="missing_path",
                                summary="Which document should I read?")
        path = _expand(str(raw))
        # Reading is not a write, but the guardrail also protects paths that
        # should not be touched at all.
        guardrails.check_path(path, writing=False)
        if not path.exists():
            return ActionResult(ok=False, error="not_found",
                                summary=f"No such file: {path}")
        if path.is_dir():
            return ActionResult(ok=False, error="is_a_directory",
                                summary=f"{path} is a folder, not a document")
        start = time.perf_counter()
        try:
            result = handler(path, args)
        except Exception as exc:      # noqa: BLE001 - malformed files are common
            result = ActionResult(
                ok=False, error=type(exc).__name__,
                summary=f"Could not read {path.name}: {exc}")
        result.duration_ms = int((time.perf_counter() - start) * 1000)
        return result

    def verify(self, ability_id: str, args: dict[str, Any],
               result: ActionResult) -> VerificationResult:
        if ability_id == "read_document":
            chars = result.evidence.get("chars", 0)
            return VerificationResult(
                verified=result.ok and chars > 0,
                strategy="text_extracted",
                detail=f"{chars} characters extracted"
                       + (" (truncated)" if result.evidence.get("truncated")
                          else ""),
                checked={"chars": chars})
        return VerificationResult(verified=result.ok, strategy="result_only",
                                  detail="observation-only ability", checked={})

    # -- handlers -----------------------------------------------------------

    def _info(self, path: Path, args: dict[str, Any]) -> ActionResult:
        stat = path.stat()
        return ActionResult(
            ok=True,
            summary=(f"{path.name}: {stat.st_size / 1024:.1f} KB, "
                     f"{path.suffix or 'no extension'}, "
                     f"{'readable' if path.suffix.lower() in self.formats() else 'unsupported format'}"),
            evidence={"path": str(path), "bytes": stat.st_size,
                      "suffix": path.suffix.lower(),
                      "supported": path.suffix.lower() in self.formats(),
                      "modified": stat.st_mtime})

    def _read(self, path: Path, args: dict[str, Any]) -> ActionResult:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return self._read_pdf(path, args)
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix in (".xlsx", ".xlsm"):
            return self._read_xlsx(path, args)
        if suffix in (".csv", ".tsv"):
            return self._read_csv(path, suffix)
        if suffix in TEXT_SUFFIXES or not suffix:
            return self._read_text(path)
        return ActionResult(
            ok=False, error="unsupported_format",
            summary=f"I cannot read {suffix} files. I can read: "
                    f"{', '.join(sorted(self.formats()))}")

    def _read_pdf(self, path: Path, args: dict[str, Any]) -> ActionResult:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            # pypdf otherwise raises something opaque; say what is actually
            # wrong so the user can supply the password or decrypt it.
            try:
                reader.decrypt("")
            except Exception:         # noqa: BLE001
                return ActionResult(
                    ok=False, error="encrypted",
                    summary=f"{path.name} is password-protected.")
        pages = len(reader.pages)
        wanted = args.get("pages")
        indices = range(pages)
        if wanted:
            indices = _page_range(str(wanted), pages)
        chunks = []
        for i in indices:
            try:
                chunks.append(reader.pages[i].extract_text() or "")
            except Exception:         # noqa: BLE001 - one bad page must not fail all
                continue
        text, truncated = _truncate("\n\n".join(c for c in chunks if c.strip()))
        if not text.strip():
            return ActionResult(
                ok=False, error="no_text",
                summary=f"{path.name} has {pages} page(s) but no extractable "
                        f"text - it is most likely a scan. OCR is not "
                        f"implemented.",
                evidence={"pages": pages, "chars": 0})
        return ActionResult(
            ok=True,
            summary=f"{path.name}: {pages} page(s), {len(text)} characters"
                    + (" (truncated)" if truncated else ""),
            evidence={"text": text, "pages": pages, "chars": len(text),
                      "truncated": truncated})

    def _read_docx(self, path: Path) -> ActionResult:
        import docx
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Tables carry the content in many real documents, so skipping them
        # loses exactly the numbers a user asks about.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text, truncated = _truncate("\n".join(parts))
        return ActionResult(
            ok=bool(text.strip()),
            summary=(f"{path.name}: {len(document.paragraphs)} paragraph(s), "
                     f"{len(document.tables)} table(s), {len(text)} characters"
                     + (" (truncated)" if truncated else ""))
            if text.strip() else f"{path.name} appears to be empty",
            evidence={"text": text, "chars": len(text), "truncated": truncated,
                      "paragraphs": len(document.paragraphs),
                      "tables": len(document.tables)})

    def _read_xlsx(self, path: Path, args: dict[str, Any]) -> ActionResult:
        import openpyxl
        # data_only: the user means the numbers, not "=SUM(B2:B40)".
        book = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
        wanted = str(args.get("sheet") or "").strip()
        names = list(book.sheetnames)
        sheets = ([s for s in names if s.lower() == wanted.lower()] or names[:1]
                  if wanted else names[:3])
        lines: list[str] = []
        empties = 0
        for name in sheets:
            ws = book[name]
            lines.append(f"--- sheet: {name} ---")
            for r, row in enumerate(ws.iter_rows(values_only=True)):
                if r >= MAX_ROWS:
                    lines.append(f"... more rows not shown (limit {MAX_ROWS})")
                    break
                if row is None or all(v is None for v in row):
                    empties += 1
                    continue
                lines.append(" | ".join("" if v is None else str(v) for v in row))
        book.close()
        text, truncated = _truncate("\n".join(lines))
        # A sheet of pure None usually means formulas that were never
        # calculated, which is worth saying instead of returning blanks.
        only_headers = len([l for l in lines if not l.startswith("---")]) == 0
        if only_headers:
            return ActionResult(
                ok=False, error="no_values",
                summary=f"{path.name} has no cached cell values. That happens "
                        f"when a file has never been opened and saved in Excel, "
                        f"so its formulas have no stored result.",
                evidence={"sheets": names, "chars": 0})
        return ActionResult(
            ok=True,
            summary=f"{path.name}: {len(names)} sheet(s) "
                    f"({', '.join(names[:4])}), {len(text)} characters"
                    + (" (truncated)" if truncated else ""),
            evidence={"text": text, "sheets": names, "chars": len(text),
                      "truncated": truncated})

    def _read_csv(self, path: Path, suffix: str) -> ActionResult:
        delimiter = "\t" if suffix == ".tsv" else ","
        rows: list[str] = []
        with path.open("r", encoding="utf-8", errors="replace", newline="") as fh:
            for i, row in enumerate(csv.reader(fh, delimiter=delimiter)):
                if i >= MAX_ROWS:
                    rows.append(f"... more rows not shown (limit {MAX_ROWS})")
                    break
                rows.append(" | ".join(row))
        text, truncated = _truncate("\n".join(rows))
        return ActionResult(
            ok=bool(text.strip()),
            summary=f"{path.name}: {len(rows)} row(s), {len(text)} characters"
                    + (" (truncated)" if truncated else ""),
            evidence={"text": text, "rows": len(rows), "chars": len(text),
                      "truncated": truncated})

    def _read_text(self, path: Path) -> ActionResult:
        raw = path.read_text(encoding="utf-8", errors="replace")
        text, truncated = _truncate(raw)
        return ActionResult(
            ok=bool(text.strip()),
            summary=f"{path.name}: {len(text)} characters"
                    + (" (truncated)" if truncated else ""),
            evidence={"text": text, "chars": len(text),
                      "truncated": truncated})


def _page_range(spec: str, total: int) -> list[int]:
    """Parse "1-5", "3", "2,4,7" into zero-based page indices."""
    out: list[int] = []
    for part in spec.replace(" ", "").split(","):
        if "-" in part:
            a, _, b = part.partition("-")
            try:
                start, end = int(a), int(b)
            except ValueError:
                continue
            out.extend(range(max(0, start - 1), min(total, end)))
        else:
            try:
                n = int(part)
            except ValueError:
                continue
            if 1 <= n <= total:
                out.append(n - 1)
    return sorted(set(out)) or list(range(total))
